"""End-to-end smoke test through the real HTTP layer (FastAPI TestClient, SQLite).

Exercises the full co-pilot loop: upload -> add job -> match -> tailor preview -> prepare
(renders the tailored DOCX in-place) -> approve (vendor stub) -> inbound interview email -> prep.

Run:  .venv\\Scripts\\python.exe scripts\\smoke_e2e.py
"""
from __future__ import annotations

import io
import os
import sys
import tempfile
from pathlib import Path

# Use an isolated temp DB/storage so the smoke run never touches dev data.
os.environ["STORAGE_DIR"] = tempfile.mkdtemp(prefix="applycopilot_smoke_")

from docx import Document  # noqa: E402
from fastapi.testclient import TestClient  # noqa: E402

from app.main import app  # noqa: E402

_results: list[bool] = []


def check(label: str, cond: bool, detail: str = "") -> None:
    _results.append(bool(cond))
    print(f"  [{'PASS' if cond else 'FAIL'}] {label}{(' — ' + detail) if detail else ''}")


def _build_resume_bytes() -> bytes:
    doc = Document()
    doc.add_paragraph("Jane Developer")
    doc.add_paragraph("jane@example.com | +1 415 555 0100 | https://github.com/jane")
    doc.add_paragraph("Summary")
    doc.add_paragraph("Backend engineer with 6 years building APIs and data pipelines.")
    doc.add_paragraph("Skills")
    doc.add_paragraph("Java, Python, AWS, Docker, PostgreSQL, React")
    doc.add_paragraph("Experience")
    doc.add_paragraph("Senior Engineer, Acme Corp, 2020-Present")
    doc.add_paragraph("Built a Python microservice on AWS serving 2M requests/day")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def main() -> int:
    # Context manager runs the lifespan (creates tables + seeds the dev tenant).
    with TestClient(app) as client:
        print("1) health")
        r = client.get("/health")
        check("health ok", r.status_code == 200 and r.json()["status"] == "ok",
              r.json().get("database"))

        print("2) upload resume")
        files = {"file": ("resume.docx", _build_resume_bytes(),
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        r = client.post("/profiles/upload", files=files)
        data = r.json()
        profile_id = data.get("profile_id")
        check("upload -> profile", r.status_code == 200 and bool(profile_id),
              f"ATS score {data.get('ats_report', {}).get('score')}")
        check("parsed skills", "Python" in data["profile"]["skills"])

        print("3) add job")
        r = client.post("/jobs", json={
            "title": "Backend Engineer", "company": "Globex", "ats_vendor": "greenhouse",
            "description": "We need Python, AWS and PostgreSQL. Build microservices at scale.",
            "url": "https://example.com/job/1",
        })
        job_id = r.json().get("id")
        check("job created", r.status_code == 200 and bool(job_id))

        print("4) match")
        r = client.post("/matches", json={"profile_id": profile_id, "job_id": job_id})
        m = r.json()
        check("match scored", r.status_code == 200 and m["score"] > 0,
              f"score={m['score']} matched={m['matched_skills']}")

        print("5) tailor preview (diff, no file yet)")
        r = client.post("/tailor/preview", json={"profile_id": profile_id, "job_id": job_id})
        t = r.json()
        check("edits produced", r.status_code == 200 and t["edit_count"] >= 1,
              f"{t['edit_count']} edits, truthful={t['truthful']}")

        print("6) prepare application (renders tailored DOCX in-place)")
        r = client.post("/applications/prepare", json={"profile_id": profile_id, "job_id": job_id})
        a = r.json()
        app_id = a.get("application_id")
        doc_path = a.get("tailored_doc_path")
        check("draft pending approval", a.get("state") == "pending_approval")
        check("tailored doc rendered", bool(doc_path) and Path(doc_path).exists(), doc_path or "")

        print("7) download tailored document")
        r = client.get(f"/applications/{app_id}/document")
        check("doc downloadable", r.status_code == 200 and len(r.content) > 0,
              f"{len(r.content)} bytes")

        print("8) approve -> submit (vendor stub -> FAILED with honest note)")
        r = client.post(f"/applications/{app_id}/approve", json={"confirm": True})
        s = r.json()
        check("approval recorded + flow ran",
              r.status_code == 200 and s["state"] in ("failed", "submitted"), s["state"])

        print("9) inbound interview email -> auto prep")
        r = client.post("/inbox/inbound", json={
            "from_addr": "recruiter@globex.com",
            "subject": "Interview invite",
            "body": "We'd love to schedule a call. Please share your availability via Calendly.",
            "application_id": app_id,
        })
        e = r.json()
        check("classified interview + prep",
              e["classification"]["category"] == "interview_invite" and "prep_plan" in e,
              ", ".join(e.get("prep_plan", {}).get("topics", [])[:4]))

    passed = all(_results)
    print("\n" + (f"ALL {len(_results)} SMOKE CHECKS PASSED" if passed
                  else f"{_results.count(False)}/{len(_results)} CHECKS FAILED"))
    return 0 if passed else 1


if __name__ == "__main__":
    sys.exit(main())
