"""Validate the full apply engine against a REAL Greenhouse form using DRY-RUN mode.

Drives a live posting end-to-end — navigate, fill identity, upload a (dummy) resume, locate the
submit button — and STOPS before submit. Uses fake demo data + a throwaway file, so nothing is
ever sent to any company. Proves the real engine (not just selectors) works on a live site.
"""
from __future__ import annotations

from pathlib import Path
from urllib.parse import urlparse

from docx import Document

from app.services.apply import playwright_greenhouse as pg
from app.services.apply.playwright_common import host_allowed
from app.services.discovery import greenhouse

STORAGE = Path(__file__).resolve().parent.parent / "storage"
DUMMY = STORAGE / "_dummy_resume.docx"
BOARDS = ["gitlab", "figma", "discord", "brex", "ramp", "airtable", "mongodb", "cloudflare"]


def _dummy_resume() -> str:
    doc = Document()
    doc.add_paragraph("Demo Applicant")
    doc.add_paragraph("demo@example.invalid | +10000000000")
    doc.add_paragraph("Skills")
    doc.add_paragraph("Python, TypeScript, React")
    STORAGE.mkdir(parents=True, exist_ok=True)
    doc.save(str(DUMMY))
    return str(DUMMY)


def main() -> int:
    job = None
    for board in BOARDS:
        try:
            jobs = greenhouse.fetch(board)
        except Exception:
            continue
        for j in jobs:
            if "greenhouse.io" in (urlparse(j.url).hostname or "") and host_allowed(j.url):
                job = j
                break
        if job:
            break
    if not job:
        print("No live Greenhouse board reachable right now.")
        return 1

    print(f"Live job: {job.title} @ {job.company}\n  {job.url}\n")
    result = pg.submit_application(
        job.url, _dummy_resume(),
        identity={"first_name": "Demo", "last_name": "Applicant",
                  "email": "demo@example.invalid", "phone": "+10000000000"},
        answers={}, headless=True, dry_run=True,
    )
    print("Engine dry-run result:")
    for k in ("ok", "status", "submit_ready", "message", "unfilled_questions"):
        print(f"   {k}: {result.get(k)}")
    print(f"   screenshot: {result.get('screenshot_path')}")
    print("\nNothing was submitted (dry run). Fake data + dummy file only.")
    return 0 if result.get("status") == "dry_run" else 2


if __name__ == "__main__":
    raise SystemExit(main())
