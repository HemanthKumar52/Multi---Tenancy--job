"""Test fixtures: programmatically build sample resume DOCX files."""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document


@pytest.fixture
def sample_resume(tmp_path: Path) -> Path:
    """A simple, ATS-friendly single-column resume with a formatted bullet."""
    doc = Document()
    doc.add_paragraph("Jane Developer")
    doc.add_paragraph("jane@example.com | +1 415 555 0100 | https://github.com/jane")

    doc.add_paragraph("Summary")
    doc.add_paragraph(
        "Backend engineer with 6 years building APIs and data pipelines."
    )

    doc.add_paragraph("Skills")
    doc.add_paragraph("Java, Python, AWS, Docker, PostgreSQL, React")

    doc.add_paragraph("Experience")
    doc.add_paragraph("Senior Engineer, Acme Corp, 2020-Present")
    # A bullet whose text spans two runs with different formatting.
    p = doc.add_paragraph(style=None)
    r1 = p.add_run("Led ")
    r1.bold = True
    p.add_run("a team of 4 building a Python microservice on AWS")

    doc.add_paragraph("Education")
    doc.add_paragraph("B.S. Computer Science, State University, 2018")

    out = tmp_path / "resume.docx"
    doc.save(str(out))
    return out


@pytest.fixture
def table_resume(tmp_path: Path) -> Path:
    """A resume that uses a layout table (ATS-hostile) — to exercise the tier-2 ATS check."""
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).paragraphs[0].add_run("John Doe")
    table.cell(0, 1).paragraphs[0].add_run("Skills")
    doc.add_paragraph("Experience")
    doc.add_paragraph("Did things with Python and AWS")
    out = tmp_path / "table_resume.docx"
    doc.save(str(out))
    return out
